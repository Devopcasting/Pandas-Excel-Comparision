import os
from fastapi import APIRouter, HTTPException, Request
from fastapi.templating import Jinja2Templates
from urllib.parse import unquote
from docx import Document
from difflib import SequenceMatcher

docx_router = APIRouter()
templates = Jinja2Templates(directory=r"app\templates")


def read_docx(file_path: str) -> str:
    """Read text from a docx file"""
    # doc = Document(file_path)
    # text = []
    # for paragraph in doc.paragraphs:
    #     text.append(paragraph.text)
    # return '\n'.join(text)

    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        line = para.text.strip()
        if line:
            full_text.append(line)
        # full_text.append(para.text)
    return '\n'.join(full_text)

def compare_documents(text1, text2):
    """
    Compare two texts and highlight differences.
    """
    matcher = SequenceMatcher(None, text1, text2)
    differences = matcher.get_opcodes()
    highlighted_text1 = ""
    highlighted_text2 = ""
    for tag, i1, i2, j1, j2 in differences:
        if tag == 'equal':
            highlighted_text1 += text1[i1:i2]
            highlighted_text2 += text2[j1:j2]
        elif tag == 'replace':
            highlighted_text1 += f"<mark>{text1[i1:i2]}</mark>"
            highlighted_text2 += f"<mark>{text2[j1:j2]}</mark>"
        elif tag == 'delete':
            highlighted_text1 += f"<mark>{text1[i1:i2]}</mark>"
        elif tag == 'insert':
            highlighted_text2 += f"<mark>{text2[j1:j2]}</mark>"
    return highlighted_text1, highlighted_text2

@docx_router.get("/compare_docx")
async def compare_docx(request: Request, file_url1: str, file_url2: str):
    # Reset the File path string
    path_file1 = unquote(file_url1)
    path_file2 = unquote(file_url2)

    # Check if the Docx documents are available in the path
    if not os.path.exists(path_file1.replace("%20", " ")):
        raise HTTPException(status_code=404, detail=f"File {path_file1} not found")
    if not os.path.exists(path_file2.replace("%20", " ")):
        raise HTTPException(status_code=404, detail=f"File {path_file2} not found")
    
    content1 = read_docx(path_file1)
    content2 = read_docx(path_file2)
    #highlighted_text1, highlighted_text2 = compare_documents(content1, content2)
   #Split content1 and content2 into lines
    lines_content1 = content1.split('\n')
    lines_content2 = content2.split('\n')

    # Compare content1 and content2 to find lines unique to each
    lines_only_in_content1 = [line for line in lines_content1 if line not in lines_content2]
    lines_only_in_content2 = [line for line in lines_content2 if line not in lines_content1]

    return templates.TemplateResponse("compare_docx.html",{"request": request, "content1": content1, "content2": content2, 
                                                           "title": "Contentverse Document Comparision",
                                                           "file1":path_file1, "file2": path_file2,
                                                           "lines_only_in_content1":lines_only_in_content1, 
                                                           "lines_only_in_content2": lines_only_in_content2})